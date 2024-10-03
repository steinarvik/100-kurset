import docx

docpath1 = "word_documents/panda1.docx"
docpath2 = "word_documents/panda2.docx"

doc1 = docx.Document(docpath1)
doc2 = docx.Document(docpath2)

#  paragraphs = doc1.paragraphs
text = doc2.paragraphs[0].text


