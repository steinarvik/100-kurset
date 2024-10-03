import os
import docx

# Define the folder containing the DOCX files and the output file path
folder_path = 'word_documents'
output_file = 'word_documents/merged_document.docx'

# Generate filenames (e.g., 'sample1.docx')
filenames = os.listdir(folder_path)

# Generate filepaths (e.g., 'word_documents/sample1.docx')
filepaths = [os.path.join(folder_path, filename) for filename in filenames if filename.endswith('.docx')]

# Create a new Document object for the merged document
merged_document = docx.Document()

# Iterate over all DOCX files in the specified folder
for filepath in filepaths:
    doc = docx.Document(filepath)

    # Append each paragraph from the current document to the merged document
    for para in doc.paragraphs:
        # Create a new paragraph in the merged document with the same formatting
        new_para = merged_document.add_paragraph()
        new_para._element.addprevious(para._element)  # Copy the element and insert it

# Save the merged document to the specified output file
merged_document.save(output_file)
